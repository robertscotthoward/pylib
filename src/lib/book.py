"""
Book class for reading text from various ebook and document formats.

Supported formats:
- txt: Plain text files
- pdf: PDF documents
- epub: EPUB ebooks
- lit: Microsoft Reader ebooks
- mobi: Kindle ebooks
- opf: Open Packaging Format (OPF) ebooks
- doc: Microsoft Word 97-2003 documents
- docx: Microsoft Word 2007+ documents
- rtf: Rich Text Format documents

The class extracts only text content, excluding images, formatting, and metadata.
"""

import os
import re
import shutil
from pathlib import Path


class Book:
    """
    A class to read text from various ebook and document formats.
    
    Supported formats: txt, pdf, epub, lit, mobi, opf, doc, docx, rtf
    
    Example:
        book = Book("path/to/book.epub")
        text = book.get_text()
    """
    
    def __init__(self, book_path):
        """
        Initialize the Book object.
        
        Args:
            book_path (str): Path to the book file
            
        Raises:
            FileNotFoundError: If the file does not exist
            ValueError: If the file format is not supported
        """
        self.book_path = book_path
        
        if not os.path.exists(book_path):
            raise FileNotFoundError(f"Book file not found: {book_path}")
        
        self.file_ext = Path(book_path).suffix.lower()
        self._validate_format()
    
    def _validate_format(self):
        """Validate that the file format is supported."""
        supported_formats = {'.txt', '.pdf', '.epub', '.lit', '.mobi', '.opf', '.doc', '.docx', '.rtf'}
        if self.file_ext not in supported_formats:
            raise ValueError(f"Unsupported format: {self.file_ext}. Supported formats: {supported_formats}")
    
    def get_text(self):
        """
        Extract text from the book file.
        
        Returns:
            str: The extracted text content, or None if extraction fails
        """
        try:
            if self.file_ext == '.txt':
                return self._read_txt()
            elif self.file_ext == '.pdf':
                return self._read_pdf()
            elif self.file_ext == '.epub':
                return self._read_epub()
            elif self.file_ext == '.lit':
                return self._read_lit()
            elif self.file_ext == '.mobi':
                return self._read_mobi()
            elif self.file_ext == '.opf':
                return self._read_opf()
            elif self.file_ext == '.doc':
                return self._read_doc()
            elif self.file_ext == '.docx':
                return self._read_docx()
            elif self.file_ext == '.rtf':
                return self._read_rtf()
        except Exception as e:
            print(f"Error reading {self.file_ext} file: {e}")
            return None
    
    def _read_txt(self):
        """Read plain text file."""
        try:
            with open(self.book_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Try with different encoding if UTF-8 fails
            with open(self.book_path, 'r', encoding='latin-1') as f:
                return f.read()
    
    def _read_pdf(self):
        """Read PDF file using pdfplumber."""
        try:
            import pdfplumber
            with pdfplumber.open(self.book_path) as pdf:
                text = '\n\n'.join([page.extract_text() or '' for page in pdf.pages])
            return text
        except ImportError:
            print("pdfplumber not installed. Install with: uv add pdfplumber")
            return None
    
    def _read_epub(self):
        """Read EPUB file."""
        try:
            from ebooklib import epub
            from bs4 import BeautifulSoup
            
            book = epub.read_epub(self.book_path)
            texts = []
            
            for item in book.get_items_of_type(epub.ITEM_DOCUMENT):
                content = item.get_content()
                soup = BeautifulSoup(content, 'html.parser')
                
                # Extract text with space separator to avoid breaking words
                text = soup.get_text(separator=' ', strip=True)
                
                # Clean up excessive whitespace
                text = re.sub(r'\s+', ' ', text)
                text = text.strip()
                
                if text:
                    texts.append(text)
            
            return '\n\n'.join(texts)
        except ImportError as e:
            print(f"ebooklib not installed. Install with: uv add ebooklib beautifulsoup4: {e}")
            return None
    
    def _read_lit(self):
        """
        Read LIT (Microsoft Reader) file.
        
        LIT files are essentially renamed ZIP files with HTML content.
        """
        try:
            import zipfile
            from bs4 import BeautifulSoup
            
            texts = []
            
            with zipfile.ZipFile(self.book_path, 'r') as zip_ref:
                # List all files in the ZIP
                for file_info in zip_ref.filelist:
                    # Look for HTML files
                    if file_info.filename.endswith(('.html', '.htm', '.xhtml')):
                        try:
                            content = zip_ref.read(file_info.filename).decode('utf-8', errors='ignore')
                            soup = BeautifulSoup(content, 'html.parser')
                            
                            # Remove script and style elements
                            for script in soup(['script', 'style']):
                                script.decompose()
                            
                            text = soup.get_text(separator=' ', strip=True)
                            text = re.sub(r'\s+', ' ', text)
                            text = text.strip()
                            
                            if text:
                                texts.append(text)
                        except Exception as e:
                            print(f"Error reading {file_info.filename} from LIT: {e}")
                            continue
            
            return '\n\n'.join(texts)
        except ImportError:
            print("beautifulsoup4 not installed. Install with: uv add beautifulsoup4")
            return None
    
    def _read_mobi(self):
        """
        Read MOBI (Kindle) file.
        
        MOBI files are binary ebook format used by Kindle.
        The mobi package extracts MOBI to HTML content.
        """
        try:
            from mobi import extract
            from bs4 import BeautifulSoup
            
            # Extract MOBI - returns a tuple: (temp_dir, html_file_path)
            extract_result = extract(self.book_path)
            
            if isinstance(extract_result, tuple):
                temp_dir = extract_result[0]
                html_path = extract_result[1]
            else:
                # Fallback if extract returns just a path
                html_path = extract_result
                temp_dir = None
            
            try:
                # Read the extracted HTML file
                with open(html_path, 'r', encoding='utf-8', errors='ignore') as f:
                    html_content = f.read()
                
                soup = BeautifulSoup(html_content, 'html.parser')
                
                # Remove script and style elements
                for script in soup(['script', 'style']):
                    script.decompose()
                
                # Extract text with newline separator to preserve paragraph boundaries
                text = soup.get_text(separator='\n', strip=True)
                
                # Clean up excessive whitespace while preserving paragraph breaks
                text = re.sub(r' +', ' ', text)  # Collapse multiple spaces to single space
                text = re.sub(r'\n\s*\n', '\n', text)  # Remove blank lines
                text = text.strip()
                
                return text if text else None
            finally:
                # Clean up the extracted files and temp directory
                if os.path.exists(html_path):
                    os.remove(html_path)
                if temp_dir and os.path.exists(temp_dir):
                    shutil.rmtree(temp_dir)
        except ImportError as e:
            print(f"Required libraries not installed. Install with: uv add mobi beautifulsoup4. Error: {e}")
            return None
        except Exception as e:
            print(f"Error reading MOBI file: {e}")
            return None
    
    def _read_opf(self):
        """
        Read OPF (Open Packaging Format) file.
        
        OPF is an XML-based format used for ebook metadata and structure.
        """
        try:
            from bs4 import BeautifulSoup
            import xml.etree.ElementTree as ET
            
            # First, try to parse as XML to find referenced content files
            tree = ET.parse(self.book_path)
            root = tree.getroot()
            
            # Define namespaces
            namespaces = {
                'opf': 'http://www.idpf.org/2007/opf',
                'dc': 'http://purl.org/dc/elements/1.1/'
            }
            
            texts = []
            
            # Find the manifest and spine to get content files
            manifest = root.find('.//opf:manifest', namespaces)
            spine = root.find('.//opf:spine', namespaces)
            
            if manifest is not None and spine is not None:
                # Get the directory of the OPF file
                opf_dir = os.path.dirname(self.book_path)
                
                # Get spine items (reading order)
                for itemref in spine.findall('opf:itemref', namespaces):
                    idref = itemref.get('idref')
                    
                    # Find the corresponding manifest item
                    for item in manifest.findall('opf:item', namespaces):
                        if item.get('id') == idref:
                            href = item.get('href')
                            if href:
                                content_path = os.path.join(opf_dir, href)
                                
                                # Read HTML/XHTML content files
                                if os.path.exists(content_path) and href.endswith(('.html', '.htm', '.xhtml')):
                                    try:
                                        with open(content_path, 'r', encoding='utf-8', errors='ignore') as f:
                                            content = f.read()
                                        
                                        soup = BeautifulSoup(content, 'html.parser')
                                        
                                        # Remove script and style elements
                                        for script in soup(['script', 'style']):
                                            script.decompose()
                                        
                                        text = soup.get_text(separator=' ', strip=True)
                                        text = re.sub(r'\s+', ' ', text)
                                        text = text.strip()
                                        
                                        if text:
                                            texts.append(text)
                                    except Exception as e:
                                        print(f"Error reading {href}: {e}")
                            break
            
            return '\n\n'.join(texts) if texts else None
        except ImportError:
            print("beautifulsoup4 not installed. Install with: uv add beautifulsoup4")
            return None
    
    def _read_doc(self):
        """
        Read DOC (Microsoft Word 97-2003) file.
        
        Uses python-docx library which has limited DOC support.
        For better DOC support, convert to DOCX first.
        """
        try:
            from docx import Document
            
            # python-docx has limited DOC support
            # Try to read as DOCX first (some DOC files are actually DOCX)
            try:
                doc = Document(self.book_path)
                texts = []
                
                for para in doc.paragraphs:
                    if para.text.strip():
                        texts.append(para.text)
                
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                texts.append(cell.text)
                
                return '\n\n'.join(texts)
            except Exception:
                # If that fails, try converting DOC to DOCX first
                return self._convert_doc_to_docx_and_read()
        except ImportError:
            print("python-docx not installed. Install with: uv add python-docx")
            return None
    
    def _convert_doc_to_docx_and_read(self):
        """Convert DOC to DOCX and read the DOCX file."""
        try:
            import subprocess
            import tempfile
            
            # Use LibreOffice to convert DOC to DOCX
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                tmp_path = tmp.name
            
            try:
                # Try using LibreOffice command line
                subprocess.run([
                    'libreoffice', '--headless', '--convert-to', 'docx',
                    '--outdir', os.path.dirname(tmp_path),
                    self.book_path
                ], check=True, capture_output=True)
                
                # Read the converted DOCX
                return self._read_docx_file(tmp_path)
            except (FileNotFoundError, subprocess.CalledProcessError):
                print("LibreOffice not available. Cannot convert DOC to DOCX.")
                return None
            finally:
                if os.path.exists(tmp_path):
                    os.remove(tmp_path)
        except Exception as e:
            print(f"Error converting DOC to DOCX: {e}")
            return None
    
    def _read_docx(self):
        """Read DOCX (Microsoft Word 2007+) file."""
        try:
            from docx import Document
            
            doc = Document(self.book_path)
            texts = []
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    texts.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            texts.append(cell.text)
            
            return '\n\n'.join(texts)
        except ImportError:
            print("python-docx not installed. Install with: uv add python-docx")
            return None
    
    def _read_docx_file(self, filepath):
        """Helper method to read a DOCX file from a given path."""
        try:
            from docx import Document
            
            doc = Document(filepath)
            texts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    texts.append(para.text)
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            texts.append(cell.text)
            
            return '\n\n'.join(texts)
        except Exception as e:
            print(f"Error reading DOCX: {e}")
            return None
    
    def _read_rtf(self):
        """Read RTF (Rich Text Format) file."""
        try:
            from striprtf.striprtf import rtf_to_text
            
            with open(self.book_path, 'r', encoding='utf-8', errors='ignore') as f:
                rtf_content = f.read()
            
            plain_text = rtf_to_text(rtf_content)
            return plain_text
        except ImportError:
            print("striprtf not installed. Install with: uv add striprtf")
            return None




def test1():
    files = r"""
E:\Books\AudioBooks\Educational\TTC (The Teaching Company) Great Courses\Philosophy\Giants of Philosophy\Homer\Homer's The Iliad\Homer's The Iliad.txt
E:\Books\Fiction\Authors\Mary Wollstonecraft Shelley\Frankenstein\Frankenstein - Mary Wollstonecraft Shelley.mobi
E:\Books\1001 Books You Must Read Before You Die\Aesop\Aesop's Fables - Aesop.epub
E:\Books\Fiction\Authors\Dickens, Charles\Dickens, Charles - Great Expectations.pdf
E:\Books\Fiction\Authors\Bradbury, Ray\Bradbury, Ray - Fahrenheit 451.doc
    """

    for file in files.strip().split("\n"):
        book = Book(file)
        text = book.get_text()
        if text:    
            text = text[:1000]
            print(text)
            print("-" * 100)
        else:
            print(f"No text found for {file}")
            print("-" * 100)


if __name__ == "__main__":
    test1()